"""
Document parsing service.
Handles: PDF, DOCX, TXT, MD, CSV, JSON, HTML, XML, ZIP
Security: Nmap XML, Nessus XML/CSV
"""
import re
import io
import csv
import json
import zipfile
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# Binary/garbage patterns to strip before embedding
GARBAGE_PATTERNS = [
    re.compile(r'endobj|xref|startxref|%%EOF', re.IGNORECASE),
    re.compile(r'<</[A-Za-z]+\s+\d+'),
    re.compile(r'obj\s*<<'),
    re.compile(r'stream\s*[\x00-\x1f]'),
    re.compile(r'[0-9a-fA-F]{20,}'),           # long hex strings
    re.compile(r'Producer|Creator|CreationDate|ModDate', re.IGNORECASE),
    re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]'),  # control chars
    re.compile(r'(\s*\n){4,}'),                  # excessive blank lines
]


@dataclass
class ParsedPage:
    page_number: int
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    filename: str
    file_type: str
    pages: List[ParsedPage]
    doc_metadata: Dict[str, Any] = field(default_factory=dict)
    is_security: bool = False
    total_pages: int = 0

    def __post_init__(self):
        self.total_pages = len(self.pages)


def clean_text(text: str) -> str:
    """Remove binary garbage, control chars, and noise from extracted text."""
    if not text:
        return ""
    for pattern in GARBAGE_PATTERNS:
        if pattern.pattern == r'(\s*\n){4,}':
            text = pattern.sub('\n\n', text)
        else:
            text = pattern.sub(' ', text)
    # Normalize whitespace
    text = re.sub(r' {3,}', '  ', text)
    text = text.strip()
    return text


def is_garbage_chunk(text: str) -> bool:
    """Return True if chunk is too noisy to be useful."""
    if len(text.strip()) < 50:
        return True
    # Check ratio of alphanumeric to total
    alnum = sum(c.isalnum() or c.isspace() for c in text)
    if len(text) > 0 and alnum / len(text) < 0.5:
        return True
    return False


class DocumentParser:

    def parse(self, file_path: Path, mode: str = "normal") -> ParsedDocument:
        """Parse any supported file type."""
        suffix = file_path.suffix.lower()
        filename = file_path.name

        try:
            if suffix == ".pdf":
                return self._parse_pdf(file_path, filename)
            elif suffix == ".docx":
                return self._parse_docx(file_path, filename)
            elif suffix in (".txt", ".md"):
                return self._parse_text(file_path, filename)
            elif suffix == ".csv":
                return self._parse_csv(file_path, filename)
            elif suffix == ".json":
                return self._parse_json(file_path, filename)
            elif suffix in (".html", ".htm"):
                return self._parse_html(file_path, filename)
            elif suffix == ".xml":
                return self._parse_xml(file_path, filename, mode)
            elif suffix == ".nessus":
                return self._parse_nessus(file_path, filename)
            elif suffix == ".zip":
                return self._parse_zip(file_path, filename, mode)
            else:
                return self._parse_text(file_path, filename)
        except Exception as e:
            logger.error(f"Parse error for {filename}: {e}")
            return ParsedDocument(
                filename=filename,
                file_type=suffix,
                pages=[ParsedPage(1, f"[Parse error: {str(e)[:200]}]")],
            )

    def _parse_pdf(self, path: Path, filename: str) -> ParsedDocument:
        import fitz  # pymupdf
        pages = []
        try:
            doc = fitz.open(str(path))
            for i, page in enumerate(doc):
                text = page.get_text("text")
                text = clean_text(text)
                if not is_garbage_chunk(text):
                    pages.append(ParsedPage(
                        page_number=i + 1,
                        text=text,
                        metadata={"width": page.rect.width, "height": page.rect.height}
                    ))
            doc.close()
        except Exception as e:
            logger.warning(f"PyMuPDF failed for {filename}, trying PyPDF2: {e}")
            try:
                import PyPDF2
                with open(path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for i, page in enumerate(reader.pages):
                        text = clean_text(page.extract_text() or "")
                        if not is_garbage_chunk(text):
                            pages.append(ParsedPage(i + 1, text))
            except Exception as e2:
                logger.error(f"Both PDF parsers failed for {filename}: {e2}")

        return ParsedDocument(filename=filename, file_type=".pdf", pages=pages)

    def _parse_docx(self, path: Path, filename: str) -> ParsedDocument:
        from docx import Document
        doc = Document(str(path))
        pages = []
        # Group paragraphs into logical pages (~50 paras each)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        page_size = 50
        for i in range(0, len(paragraphs), page_size):
            chunk = "\n".join(paragraphs[i:i + page_size])
            text = clean_text(chunk)
            if not is_garbage_chunk(text):
                pages.append(ParsedPage(page_number=i // page_size + 1, text=text))

        return ParsedDocument(filename=filename, file_type=".docx", pages=pages)

    def _parse_text(self, path: Path, filename: str) -> ParsedDocument:
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
        except Exception:
            text = path.read_bytes().decode("latin-1", errors="replace")

        lines = text.splitlines()
        page_size = 200
        pages = []
        for i in range(0, len(lines), page_size):
            chunk = "\n".join(lines[i:i + page_size])
            chunk = clean_text(chunk)
            if not is_garbage_chunk(chunk):
                pages.append(ParsedPage(i // page_size + 1, chunk))

        return ParsedDocument(filename=filename, file_type=path.suffix, pages=pages)

    def _parse_csv(self, path: Path, filename: str) -> ParsedDocument:
        pages = []
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                reader = csv.reader(f)
                rows = list(reader)

            headers = rows[0] if rows else []
            page_size = 100
            for i in range(1, len(rows), page_size):
                batch = rows[i:i + page_size]
                lines = []
                for row in batch:
                    if any(cell.strip() for cell in row):
                        pair = "; ".join(
                            f"{h}: {v}" for h, v in zip(headers, row) if v.strip()
                        )
                        lines.append(pair)
                text = clean_text("\n".join(lines))
                if not is_garbage_chunk(text):
                    pages.append(ParsedPage(i // page_size + 1, text))
        except Exception as e:
            logger.error(f"CSV parse error: {e}")

        return ParsedDocument(filename=filename, file_type=".csv", pages=pages)

    def _parse_json(self, path: Path, filename: str) -> ParsedDocument:
        try:
            data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
            text = json.dumps(data, indent=2, ensure_ascii=False)
        except Exception:
            text = path.read_text(encoding="utf-8", errors="replace")

        lines = text.splitlines()
        page_size = 300
        pages = []
        for i in range(0, len(lines), page_size):
            chunk = clean_text("\n".join(lines[i:i + page_size]))
            if not is_garbage_chunk(chunk):
                pages.append(ParsedPage(i // page_size + 1, chunk))

        return ParsedDocument(filename=filename, file_type=".json", pages=pages)

    def _parse_html(self, path: Path, filename: str) -> ParsedDocument:
        from bs4 import BeautifulSoup
        html = path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(html, "lxml")
        for tag in soup(["script", "style", "meta", "link"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        text = clean_text(text)
        lines = text.splitlines()
        page_size = 200
        pages = []
        for i in range(0, len(lines), page_size):
            chunk = clean_text("\n".join(lines[i:i + page_size]))
            if not is_garbage_chunk(chunk):
                pages.append(ParsedPage(i // page_size + 1, chunk))
        return ParsedDocument(filename=filename, file_type=".html", pages=pages)

    def _parse_xml(self, path: Path, filename: str, mode: str = "normal") -> ParsedDocument:
        """Handles generic XML and Nmap XML in security mode."""
        from bs4 import BeautifulSoup
        content = path.read_text(encoding="utf-8", errors="replace")

        # Detect Nmap format
        if "<nmaprun" in content.lower():
            return self._parse_nmap_xml(path, filename, content)

        # Generic XML
        soup = BeautifulSoup(content, "xml")
        text = clean_text(soup.get_text(separator="\n"))
        lines = text.splitlines()
        page_size = 200
        pages = []
        for i in range(0, len(lines), page_size):
            chunk = clean_text("\n".join(lines[i:i + page_size]))
            if not is_garbage_chunk(chunk):
                pages.append(ParsedPage(i // page_size + 1, chunk))
        return ParsedDocument(filename=filename, file_type=".xml", pages=pages)

    def _parse_nmap_xml(self, path: Path, filename: str, content: str) -> ParsedDocument:
        """Parse Nmap XML scan results into structured text."""
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(content, "xml")
        pages = []
        page_num = 1

        scan_info = soup.find("nmaprun")
        if scan_info:
            info_text = (
                f"Nmap Scan Results\n"
                f"Scanner: {scan_info.get('scanner', 'nmap')}\n"
                f"Command: {scan_info.get('args', '')}\n"
                f"Start Time: {scan_info.get('startstr', '')}\n"
            )
            pages.append(ParsedPage(page_num, info_text, {"section": "scan_header"}))
            page_num += 1

        for host in soup.find_all("host"):
            host_lines = []
            # Status
            status = host.find("status")
            if status:
                host_lines.append(f"Host Status: {status.get('state', 'unknown')}")
            # Address
            for addr in host.find_all("address"):
                host_lines.append(f"Address: {addr.get('addr', '')} ({addr.get('addrtype', '')})")
            # Hostname
            for hn in host.find_all("hostname"):
                host_lines.append(f"Hostname: {hn.get('name', '')}")
            # Ports
            for port in host.find_all("port"):
                state = port.find("state")
                service = port.find("service")
                port_line = f"Port: {port.get('portid', '')}/{port.get('protocol', '')}"
                if state:
                    port_line += f" | State: {state.get('state', '')}"
                if service:
                    port_line += (
                        f" | Service: {service.get('name', '')}"
                        f" {service.get('product', '')} {service.get('version', '')}".strip()
                    )
                    cpe = service.find("cpe")
                    if cpe:
                        port_line += f" | CPE: {cpe.get_text()}"
                host_lines.append(port_line)
            # OS detection
            for osmatch in host.find_all("osmatch"):
                host_lines.append(
                    f"OS Detection: {osmatch.get('name', '')} "
                    f"(accuracy: {osmatch.get('accuracy', '')}%)"
                )
            # Scripts (NSE)
            for script in host.find_all("script"):
                host_lines.append(f"Script [{script.get('id', '')}]: {script.get('output', '')[:300]}")

            text = "\n".join(host_lines)
            if not is_garbage_chunk(text):
                pages.append(ParsedPage(page_num, text, {"section": "host", "type": "nmap"}))
                page_num += 1

        return ParsedDocument(
            filename=filename, file_type=".xml", pages=pages, is_security=True,
            doc_metadata={"format": "nmap_xml"}
        )

    def _parse_nessus(self, path: Path, filename: str) -> ParsedDocument:
        """Parse Nessus .nessus (XML) scan output."""
        from bs4 import BeautifulSoup
        content = path.read_text(encoding="utf-8", errors="replace")
        soup = BeautifulSoup(content, "xml")
        pages = []
        page_num = 1

        for report_host in soup.find_all("ReportHost"):
            host_name = report_host.get("name", "unknown")
            for item in report_host.find_all("ReportItem"):
                lines = [
                    f"Host: {host_name}",
                    f"Plugin: {item.get('pluginName', '')}",
                    f"Plugin ID: {item.get('pluginID', '')}",
                    f"Port: {item.get('port', '')}/{item.get('protocol', '')}",
                    f"Severity: {item.get('severity', '')}",
                    f"Service: {item.get('svc_name', '')}",
                ]
                for tag in ["description", "solution", "synopsis", "risk_factor",
                            "cvss_base_score", "cvss3_base_score", "cve", "cwe",
                            "see_also", "plugin_output"]:
                    el = item.find(tag)
                    if el and el.get_text(strip=True):
                        lines.append(f"{tag.replace('_', ' ').title()}: {el.get_text(strip=True)[:500]}")

                text = "\n".join(lines)
                if not is_garbage_chunk(text):
                    pages.append(ParsedPage(page_num, text, {
                        "section": "finding",
                        "host": host_name,
                        "severity": item.get("severity", ""),
                        "plugin_id": item.get("pluginID", ""),
                    }))
                    page_num += 1

        return ParsedDocument(
            filename=filename, file_type=".nessus", pages=pages, is_security=True,
            doc_metadata={"format": "nessus"}
        )

    def _parse_zip(self, path: Path, filename: str, mode: str) -> ParsedDocument:
        """Extract and parse all supported files from a ZIP archive."""
        all_pages = []
        page_offset = 0
        try:
            with zipfile.ZipFile(str(path), "r") as zf:
                for name in zf.namelist():
                    suffix = Path(name).suffix.lower()
                    if suffix in (".pdf", ".docx", ".txt", ".md", ".csv",
                                   ".json", ".html", ".xml", ".nessus"):
                        try:
                            data = zf.read(name)
                            tmp = Path(f"/tmp/evara_zip_{Path(name).name}")
                            tmp.write_bytes(data)
                            sub_doc = self.parse(tmp, mode)
                            for page in sub_doc.pages:
                                page.page_number += page_offset
                                page.metadata["zip_entry"] = name
                                all_pages.append(page)
                            page_offset += len(sub_doc.pages)
                            tmp.unlink(missing_ok=True)
                        except Exception as e:
                            logger.warning(f"ZIP entry {name} failed: {e}")
        except Exception as e:
            logger.error(f"ZIP parse failed: {e}")

        return ParsedDocument(filename=filename, file_type=".zip", pages=all_pages)


parser = DocumentParser()
